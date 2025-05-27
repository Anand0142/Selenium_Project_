import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import google.generativeai as genai
from supabase import create_client
import requests
import time

# Configuration
SUPABASE_URL = "YOUR_SUPABASE_URL"
SUPABASE_KEY = "SUPABASE_KEY"
GEMINI_API_KEY = "YOUR_GEMINI_KEY"
MODEL_NAME = "models/gemini-1.5-pro-latest"

# JSearch API config
JSEARCH_API_KEY = "YOUR_API_KEY"
JSEARCH_HEADERS = {
    "X-RapidAPI-Key": JSEARCH_API_KEY,
    "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
}

# Initialize clients
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
genai.configure(api_key=GEMINI_API_KEY)

def fetch_job_description(job_id):
    """Fetch job description from Supabase or JSearch API"""
    try:
        # First try to get from Supabase
        job = supabase.table("jobs").select("*").eq("id", job_id).execute()
        if job.data:
            return job.data[0].get("description", "")
            
        # If not in Supabase, try JSearch API
        url = "https://jsearch.p.rapidapi.com/job-details"
        querystring = {"job_id": job_id}
        response = requests.get(url, headers=JSEARCH_HEADERS, params=querystring)
        response.raise_for_status()
        
        job_data = response.json().get("data", {})
        return job_data.get("job_description", "")
        
    except Exception as e:
        print(f"⚠️ Error fetching job description: {str(e)}")
        return ""

def preserve_structure_customize(original_doc, job_desc):
    """Customizes resume while preserving exact structure"""
    # 1. Extract original structure
    original_content = []
    for para in original_doc.paragraphs:
        original_content.append({
            'text': para.text,
            'style': para.style.name,
            'runs': [{'text': run.text, 'bold': run.bold} for run in para.runs]
        })
    
    # 2. Extract text for analysis
    full_text = "\n".join([item['text'] for item in original_content])
    
    # 3. Get customization suggestions
    prompt = f"""
    Analyze this resume against the job description.
    Return ONLY specific improvements to apply to the EXISTING structure:
    
    RESUME:
    {full_text[:15000]}
    
    JOB DESCRIPTION:
    {job_desc[:5000]}
    
    INSTRUCTIONS:
    1. Identify which existing skills to **bold** (format exactly as they appear)
    2. Suggest minor text improvements ONLY within existing sections
    3. Never add/remove sections
    4. Never change formatting except bolding keywords
    5. Return EXACT original text with your changes
    """
    
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            prompt,
            generation_config={"temperature": 0.1}
        )
        return response.text
    except Exception as e:
        print(f"⚠️ Customization failed: {str(e)}")
        return full_text  # Fallback to original

def process_resumes():
    try:
        # Fetch data
        resumes = supabase.table("resumes").select("*").execute().data
        jobs = supabase.table("jobs").select("*").execute().data
        
        for resume in resumes:
            if not resume.get('public_url'):
                continue
                
            try:
                # Download resume
                response = requests.get(resume['public_url'], timeout=10)
                file_bytes = io.BytesIO(response.content)
                
                # Load original document
                if resume['public_url'].lower().endswith('.docx'):
                    original_doc = Document(file_bytes)
                else:  # Convert PDF to DOCX
                    pdf_text = ""
                    with fitz.open(stream=file_bytes.getvalue(), filetype="pdf") as doc:
                        pdf_text = "\n".join([page.get_text() for page in doc])
                    original_doc = Document()
                    original_doc.add_paragraph(pdf_text)
                
                for job in jobs:
                    print(f"\n🔧 Enhancing resume {resume['id'][:8]} for job {job['id'][:8]}")
                    
                    # Get job description
                    job_desc = fetch_job_description(job['id'])
                    if not job_desc:
                        print(f"⚠️ No job description found for job {job['id']}")
                        continue
                    
                    # Customize while preserving structure
                    customized_text = preserve_structure_customize(original_doc, job_desc)
                    
                    # Rebuild document with original structure
                    new_doc = Document()
                    for para in original_doc.paragraphs:
                        new_para = new_doc.add_paragraph(style=para.style)
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.bold = run.bold
                    
                    # Apply simple keyword bolding (non-destructive)
                    for para in new_doc.paragraphs:
                        if 'skills' in para.text.lower():
                            for run in para.runs:
                                if any(kw.lower() in run.text.lower() 
                                      for kw in job_desc.split()[:20]):  # Simple keyword matching
                                    run.bold = True
                    
                    # Upload
                    output = io.BytesIO()
                    new_doc.save(output)
                    supabase.storage.from_("customizedresumes").upload(
                        f"enhanced_{resume['id']}_{job['id']}.docx",
                        output.getvalue()
                    )
                    print(f"✅ Preserved-structure resume uploaded")
                    
                    time.sleep(1.1)  # Rate limiting
            
            except Exception as e:
                print(f"⚠️ Failed to process resume {resume['id']}: {str(e)}")
                continue
                
    except Exception as e:
        print(f"💥 Fatal error: {str(e)}")

if __name__ == "__main__":
    process_resumes()
